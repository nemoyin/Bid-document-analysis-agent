"""
报告生成器
从分析结果生成 PDF 和 Word 格式的投标分析报告（6维度）。
"""

from __future__ import annotations

import json
import os
import uuid
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any

from loguru import logger
from sqlalchemy import select

from app.core.config import settings
from app.models.analysis import (
    AnalysisTask,
    ErrorDetectionResult,
    ImageSimilarityResult,
    SimilarityResult,
)
from app.models.project import Project


class ReportGenerator:
    """分析报告生成器。"""

    REPORT_DIR = Path("./data/reports")

    def __init__(self):
        self.REPORT_DIR.mkdir(parents=True, exist_ok=True)

    async def _collect_report_data(
        self, db_session_factory, project_id: uuid.UUID, task_id: uuid.UUID
    ) -> dict[str, Any]:
        async with db_session_factory() as db:
            proj_result = await db.execute(
                select(Project).where(Project.id == project_id)
            )
            project = proj_result.scalar_one_or_none()
            task_result = await db.execute(
                select(AnalysisTask).where(AnalysisTask.id == task_id)
            )
            task = task_result.scalar_one_or_none()
            sim_result = await db.execute(
                select(SimilarityResult).where(SimilarityResult.task_id == task_id)
            )
            sim_results = sim_result.scalars().all()
            err_result = await db.execute(
                select(ErrorDetectionResult).where(
                    ErrorDetectionResult.task_id == task_id
                )
            )
            err_results = err_result.scalars().all()
            img_result = await db.execute(
                select(ImageSimilarityResult).where(
                    ImageSimilarityResult.task_id == task_id
                )
            )
            img_results = img_result.scalars().all()

            # 解析6维度评分
            dims = {}
            if task and task.error_message:
                try:
                    parsed = json.loads(task.error_message)
                    if isinstance(parsed, dict) and "text_score" in parsed:
                        dims = parsed
                except (ValueError, TypeError):
                    pass

            return {
                "project": project,
                "task": task,
                "sim_results": sim_results,
                "err_results": err_results,
                "img_results": img_results,
                "dimension_scores": dims,
            }

    async def generate_report(
        self,
        db_session_factory,
        project_id: uuid.UUID,
        task_id: uuid.UUID,
        output_format: str = "pdf",
    ) -> tuple[bytes, str]:
        data = await self._collect_report_data(db_session_factory, project_id, task_id)
        project = data["project"]
        task = data["task"]

        if not project or not task:
            raise ValueError("项目或分析任务不存在")

        if output_format == "pdf":
            file_bytes, filename = self._generate_pdf(data)
        elif output_format in ("word", "docx"):
            file_bytes, filename = self._generate_word(data)
        else:
            raise ValueError(f"不支持的格式: {output_format}")

        report_id = uuid.uuid4()
        report_path = self.REPORT_DIR / str(project_id)
        report_path.mkdir(parents=True, exist_ok=True)
        report_file = report_path / f"{report_id}.{output_format}"
        report_file.write_bytes(file_bytes)
        logger.info(f"报告已生成: {report_file}")
        return file_bytes, filename

    @staticmethod
    def _risk_level_cn(level: str | None) -> str:
        mapping = {"low": "低风险", "moderate": "中风险", "high": "高风险", "critical": "严重风险"}
        return mapping.get((level or "").lower(), level or "未知")

    # ═══════════════════════════════════════════════════════════
    # PDF
    # ═══════════════════════════════════════════════════════════

    def _generate_pdf(self, data: dict) -> tuple[bytes, str]:
        project = data["project"]
        task = data["task"]
        dims = data.get("dimension_scores", {})

        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
            from reportlab.lib.units import mm
            from reportlab.platypus import (
                PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle,
            )
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.lib import colors

            # 注册中文字体
            font_name = "Helvetica"
            for fp in [
                "C:/Windows/Fonts/simhei.ttf",
                "C:/Windows/Fonts/msyh.ttc",
                "C:/Windows/Fonts/simsun.ttc",
            ]:
                if os.path.exists(fp):
                    try:
                        pdfmetrics.registerFont(TTFont("ChineseFont", fp))
                        font_name = "ChineseFont"
                        break
                    except Exception:
                        continue

            buf = BytesIO()
            doc = SimpleDocTemplate(
                buf, pagesize=A4,
                rightMargin=20*mm, leftMargin=20*mm,
                topMargin=20*mm, bottomMargin=20*mm,
                title=f"{project.name} 分析报告",
            )

            styles = getSampleStyleSheet()
            body = ParagraphStyle("cn", parent=styles["Normal"],
                fontName=font_name, fontSize=10, leading=16, spaceAfter=6)
            title_s = ParagraphStyle("cnt", parent=styles["Title"],
                fontName=font_name, fontSize=20, leading=28, alignment=1)
            h2 = ParagraphStyle("cnh2", parent=styles["Heading2"],
                fontName=font_name, fontSize=13, leading=20, spaceBefore=14, spaceAfter=6)
            h3 = ParagraphStyle("cnh3", parent=styles["Heading3"],
                fontName=font_name, fontSize=11, leading=16, spaceBefore=10, spaceAfter=4)

            story = []

            # ── 封面 ──
            story.append(Spacer(1, 50))
            story.append(Paragraph("投标标书智能分析报告", title_s))
            story.append(Spacer(1, 20))
            story.append(Paragraph(f"项目名称：{project.name}", body))
            date_str = task.completed_at.strftime("%Y-%m-%d %H:%M") if task.completed_at else "-"
            story.append(Paragraph(f"分析时间：{date_str}", body))
            story.append(Paragraph(f"风险等级：{self._risk_level_cn(task.risk_level)}", body))
            story.append(Paragraph(f"综合评分：{float(task.risk_score or 0):.1f} 分", body))
            story.append(PageBreak())

            # ── 一、分析摘要 ──
            story.append(Paragraph("一、分析摘要", h2))
            story.append(Paragraph(
                f"本项目共有 {project.file_count} 个投标文件参与分析。"
                f"综合风险评分 {float(task.risk_score or 0):.1f} 分，"
                f"风险等级为 {self._risk_level_cn(task.risk_level)}。",
                body))

            # ── 二、6维度评分总览 ──
            story.append(Spacer(1, 8))
            story.append(Paragraph("二、各维度评分总览", h2))
            dim_headers = ["维度", "满分", "得分", "占比"]
            dim_rows = [dim_headers]
            dim_defs = [
                ("文本相似度", 30, dims.get("text_score", 0)),
                ("目录结构相似", 15, dims.get("structure_score", 0)),
                ("图片相似度", 15, dims.get("image_score", 0)),
                ("表格相似度", 10, dims.get("table_score", 0)),
                ("错别字一致性", 20, dims.get("error_score", 0)),
                ("元数据一致性", 10, dims.get("metadata_score", 0)),
            ]
            for label, full, score in dim_defs:
                pct = score * 100
                dim_rows.append([label, str(full), f"{pct:.1f}", f"{pct * full / 100:.1f}"])

            dim_table = Table(dim_rows, colWidths=[110, 50, 60, 50])
            dim_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("FONTNAME", (0, 0), (-1, -1), font_name),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ALIGN", (1, 0), (-1, -1), "CENTER"),
            ]))
            story.append(dim_table)
            story.append(Spacer(1, 6))
            story.append(Paragraph(
                f"综合评分 = 文本({dim_defs[0][2]*dim_defs[0][1]:.1f}) + "
                f"结构({dim_defs[1][2]*dim_defs[1][1]:.1f}) + "
                f"图片({dim_defs[2][2]*dim_defs[2][1]:.1f}) + "
                f"表格({dim_defs[3][2]*dim_defs[3][1]:.1f}) + "
                f"错误({dim_defs[4][2]*dim_defs[4][1]:.1f}) + "
                f"元数据({dim_defs[5][2]*dim_defs[5][1]:.1f})",
                body))

            # ── 三、文档相似度明细 ──
            story.append(PageBreak())
            story.append(Paragraph("三、文档相似度明细", h2))
            sims = data["sim_results"]
            story.append(Paragraph(f"共发现 {len(sims)} 对相似文档。", body))

            if sims:
                sim_headers = ["文档A", "文档B", "全文(%)", "结构(%)", "表格(%)", "元数据(%)"]
                sim_rows = [sim_headers]
                for s in sorted(sims, key=lambda x: float(x.full_text_similarity or 0), reverse=True)[:10]:
                    sim_rows.append([
                        str(s.doc1_id)[:8],
                        str(s.doc2_id)[:8],
                        f"{float(s.full_text_similarity or 0):.1f}",
                        f"{float(s.structure_similarity or 0):.1f}",
                        f"{float(s.table_similarity or 0):.1f}",
                        f"{float(s.metadata_consistency or 0):.1f}",
                    ])

                sim_table = Table(sim_rows, colWidths=[70, 70, 60, 60, 60, 70])
                sim_table.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTSIZE", (0, 0), (-1, -1), 8),
                    ("FONTNAME", (0, 0), (-1, -1), font_name),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("ALIGN", (2, 0), (-1, -1), "CENTER"),
                ]))
                story.append(sim_table)

            # ── 四、图片相似度 ──
            story.append(Spacer(1, 10))
            story.append(Paragraph("四、图片相似度分析", h2))
            imgs = data["img_results"]
            story.append(Paragraph(f"共发现 {len(imgs)} 对相似图片。", body))
            if imgs:
                for img in imgs[:10]:
                    story.append(Paragraph(
                        f"  • 相似度 {float(img.similarity_score or 0):.1f}% "
                        f"(算法: {img.hash_algorithm})",
                        body))

            # ── 五、元数据一致性 ──
            story.append(Spacer(1, 6))
            story.append(Paragraph("五、元数据一致性分析", h2))
            meta_matches: list[str] = []
            for s in sims:
                if s.details and isinstance(s.details, dict):
                    mc = s.details.get("metadata_comparison", {})
                    for field in mc.get("matched_fields", []):
                        if field not in meta_matches:
                            meta_matches.append(field)
            if meta_matches:
                labels = {"author": "作者", "creator": "创建者", "producer": "生成软件",
                          "title": "标题", "company": "公司", "last_modified_by": "修改者"}
                match_text = "、".join(labels.get(f, f) for f in meta_matches)
                story.append(Paragraph(
                    f"不同企业的投标文件在以下元数据字段上存在一致：{match_text}。"
                    f"这可能表明这些标书出自同一来源，建议重点关注。", body))
            else:
                story.append(Paragraph("未发现元数据异常一致性。", body))

            # ── 六、错误检测 ──
            story.append(Spacer(1, 6))
            story.append(Paragraph("六、错误与一致性分析", h2))
            errs = data["err_results"]
            story.append(Paragraph(f"共检测到 {len(errs)} 条问题。", body))
            if errs:
                shared_count = sum(1 for e in errs if e.is_shared)
                story.append(Paragraph(f"其中跨文档共享错误 {shared_count} 条。", body))
                for err in errs[:20]:
                    story.append(Paragraph(
                        f"[{err.error_type}] {err.original_text[:80]}"
                        + (f" → {err.corrected_text}" if err.corrected_text else ""),
                        body))

            # ── 七、综合建议 ──
            story.append(Spacer(1, 10))
            story.append(Paragraph("七、综合建议", h2))
            rl = (task.risk_level or "LOW").upper()
            suggestions_map = {
                "CRITICAL": [
                    "1. 投标文件间存在严重相似度，强烈建议启动围标串标调查程序。",
                    "2. 重点关注高相似文本段落、同源图片和共享错误。",
                    "3. 元数据一致性异常强烈暗示标书可能出自同一源头。",
                ],
                "HIGH": [
                    "1. 存在明显相似度，建议人工复核是否存在围标嫌疑。",
                    "2. 重点核查高相似片段和共享错误一致性问题。",
                ],
                "MEDIUM": [
                    "1. 存在一定程度的相似度，建议抽查确认合理性。",
                ],
                "LOW": [
                    "1. 未发现明显异常，建议归档备查。",
                ],
            }
            for s in suggestions_map.get(rl, suggestions_map["LOW"]):
                story.append(Paragraph(s, body))

            doc.build(story)
            buf.seek(0)
            return buf.read(), self._make_filename(project, "pdf")

        except ImportError:
            logger.warning("reportlab 未安装，PDF 降级为文本")
            return self._gen_text(data), self._make_filename(project, "pdf")

    # ═══════════════════════════════════════════════════════════
    # Word
    # ═══════════════════════════════════════════════════════════

    def _generate_word(self, data: dict) -> tuple[bytes, str]:
        project = data["project"]
        task = data["task"]
        dims = data.get("dimension_scores", {})
        sims = data["sim_results"]
        imgs = data["img_results"]
        errs = data["err_results"]

        try:
            from docx import Document
            from docx.shared import Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT

            doc = Document()
            style = doc.styles["Normal"]
            style.font.size = Pt(11)

            # 封面
            doc.add_paragraph()
            t = doc.add_paragraph()
            t.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = t.add_run("投标标书智能分析报告")
            r.font.size = Pt(22)
            r.bold = True
            doc.add_paragraph()
            info = doc.add_paragraph()
            info.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_str = task.completed_at.strftime("%Y-%m-%d %H:%M") if task.completed_at else "-"
            info.add_run(f"项目：{project.name}\n分析时间：{date_str}\n"
                         f"风险等级：{self._risk_level_cn(task.risk_level)}\n"
                         f"综合评分：{float(task.risk_score or 0):.1f} 分")
            doc.add_page_break()

            # 摘要
            doc.add_heading("一、分析摘要", level=1)
            doc.add_paragraph(
                f"本项目共 {project.file_count} 个投标文件。"
                f"综合风险评分 {float(task.risk_score or 0):.1f} 分，"
                f"等级 {self._risk_level_cn(task.risk_level)}。")

            # 6维度评分
            doc.add_heading("二、各维度评分总览", level=1)
            dim_defs = [
                ("文本相似度", 30, dims.get("text_score", 0)),
                ("目录结构相似", 15, dims.get("structure_score", 0)),
                ("图片相似度", 15, dims.get("image_score", 0)),
                ("表格相似度", 10, dims.get("table_score", 0)),
                ("错别字一致性", 20, dims.get("error_score", 0)),
                ("元数据一致性", 10, dims.get("metadata_score", 0)),
            ]
            dim_table = doc.add_table(rows=1+len(dim_defs), cols=4)
            dim_table.style = "Light Grid Accent 1"
            dim_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for ci, hdr in enumerate(["维度", "满分", "得分", "占比"]):
                dim_table.rows[0].cells[ci].text = hdr
            for ri, (label, full, score) in enumerate(dim_defs):
                cells = dim_table.rows[ri+1].cells
                cells[0].text = label
                cells[1].text = str(full)
                cells[2].text = f"{score*100:.1f}"
                cells[3].text = f"{score*full:.1f}"

            # 相似度明细
            doc.add_heading("三、文档相似度明细", level=1)
            doc.add_paragraph(f"共发现 {len(sims)} 对相似文档。")
            if sims:
                st = doc.add_table(rows=1, cols=6)
                st.style = "Light Grid Accent 1"
                for ci, hdr in enumerate(["文档A", "文档B", "全文%", "结构%", "表格%", "元数据%"]):
                    st.rows[0].cells[ci].text = hdr
                for s in sorted(sims, key=lambda x: float(x.full_text_similarity or 0), reverse=True)[:10]:
                    row = st.add_row().cells
                    row[0].text = str(s.doc1_id)[:8]
                    row[1].text = str(s.doc2_id)[:8]
                    row[2].text = f"{float(s.full_text_similarity or 0):.1f}"
                    row[3].text = f"{float(s.structure_similarity or 0):.1f}"
                    row[4].text = f"{float(s.table_similarity or 0):.1f}"
                    row[5].text = f"{float(s.metadata_consistency or 0):.1f}"

            # 图片
            doc.add_heading("四、图片相似度分析", level=1)
            doc.add_paragraph(f"共发现 {len(imgs)} 对相似图片。")

            # 元数据
            doc.add_heading("五、元数据一致性分析", level=1)
            meta_fields: list[str] = []
            for s in sims:
                if s.details and isinstance(s.details, dict):
                    for f in s.details.get("metadata_comparison", {}).get("matched_fields", []):
                        if f not in meta_fields:
                            meta_fields.append(f)
            labels = {"author": "作者", "creator": "创建者", "producer": "生成软件",
                      "title": "标题", "company": "公司", "last_modified_by": "修改者"}
            if meta_fields:
                doc.add_paragraph(f"匹配字段：{'、'.join(labels.get(f, f) for f in meta_fields)}")
            else:
                doc.add_paragraph("未发现元数据异常一致性。")

            # 错误
            doc.add_heading("六、错误与一致性分析", level=1)
            doc.add_paragraph(f"共检测到 {len(errs)} 条问题。")
            shared = sum(1 for e in errs if e.is_shared)
            doc.add_paragraph(f"其中跨文档共享错误 {shared} 条。")

            # 建议
            doc.add_heading("七、综合建议", level=1)
            rl = (task.risk_level or "LOW").upper()
            sug = {
                "CRITICAL": ["强烈建议启动围标串标调查", "重点关注高相似段落和共享错误"],
                "HIGH": ["建议人工复核是否存在围标嫌疑"],
                "MEDIUM": ["建议抽查确认合理性"],
                "LOW": ["未发现明显异常，建议归档备查"],
            }
            for s in sug.get(rl, sug["LOW"]):
                doc.add_paragraph(s, style="List Number")

            buf = BytesIO()
            doc.save(buf)
            buf.seek(0)
            return buf.read(), self._make_filename(project, "docx")

        except ImportError:
            logger.warning("python-docx 未安装")
            return self._gen_text(data), self._make_filename(project, "docx")

    # ═══════════════════════════════════════════════════════════
    # 降级
    # ═══════════════════════════════════════════════════════════

    def _gen_text(self, data: dict) -> bytes:
        project = data["project"]
        task = data["task"]
        dims = data.get("dimension_scores", {})
        lines = [
            "=" * 60, "  投标标书智能分析报告", "=" * 60,
            f"项目: {project.name}",
            f"综合评分: {float(task.risk_score or 0):.1f} 分",
            f"风险等级: {self._risk_level_cn(task.risk_level)}",
            "",
            "各维度评分:",
        ]
        for label, full, key in [
            ("文本相似度", 30, "text_score"),
            ("目录结构", 15, "structure_score"),
            ("图片相似度", 15, "image_score"),
            ("表格相似度", 10, "table_score"),
            ("错别字一致性", 20, "error_score"),
            ("元数据一致性", 10, "metadata_score"),
        ]:
            s = dims.get(key, 0) * 100
            lines.append(f"  {label} ({full}分): {s:.1f}分")
        lines += [
            "",
            f"相似文档对: {len(data['sim_results'])} 对",
            f"相似图片: {len(data['img_results'])} 对",
            f"错误检测: {len(data['err_results'])} 条",
        ]
        return "\n".join(lines).encode("utf-8")

    @staticmethod
    def _make_filename(project: Project, ext: str) -> str:
        date_str = datetime.now().strftime("%Y%m%d")
        safe = project.name.replace("/", "_").replace("\\", "_")[:50]
        return f"{safe}_分析报告_{date_str}.{ext}"
