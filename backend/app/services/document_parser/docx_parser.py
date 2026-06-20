"""
DOCX 文档解析器。
使用 python-docx 提取段落、表格和图片。
"""

from __future__ import annotations

import hashlib
import io
import os
from typing import Optional

from loguru import logger

from app.services.document_parser.base import (
    BaseDocumentParser,
    ImageContent,
    PageContent,
    ParseResult,
)


class DocxParser(BaseDocumentParser):
    """DOCX 文档解析器。

    使用 python-docx 库解析 Word 文档。
    DOCX 没有显式的分页概念，因此按段落分块模拟分页（每页约 40 个段落）。
    """

    PARAGRAPHS_PER_PAGE = 40

    def __init__(self, extract_images: bool = True):
        """初始化 DOCX 解析器。

        Args:
            extract_images: 是否提取图片
        """
        self.extract_images = extract_images
        self._docx = None
        self._lazy_import()

    def _lazy_import(self) -> None:
        """惰性导入 python-docx 库。"""
        try:
            import docx
            self._docx = docx
        except ImportError:
            logger.warning("python-docx 未安装，DOCX 解析将不可用")

    def parse(self, file_path: str) -> ParseResult:
        """解析 DOCX 文档。

        Args:
            file_path: DOCX 文件路径

        Returns:
            ParseResult: 解析结果

        Raises:
            FileNotFoundError: 文件不存在
            ValueError: 解析失败
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"DOCX 文件不存在: {file_path}")

        if not self._docx:
            raise ValueError("python-docx 未安装，无法解析 DOCX 文件")

        file_name = os.path.basename(file_path)
        return self._parse_docx(file_path, file_name)

    def _parse_docx(self, file_path: str, file_name: str) -> ParseResult:
        """执行 DOCX 解析。"""
        from docx import Document as DocxDocument
        from docx.opc.constants import RELATIONSHIP_TYPE as RT

        doc = DocxDocument(file_path)

        pages: list[PageContent] = []
        images: list[ImageContent] = []

        # 提取元数据
        metadata = {
            "file_name": file_name,
            "file_type": "docx",
            "title": doc.core_properties.title or "",
            "author": doc.core_properties.author or "",
            "created": str(doc.core_properties.created) if doc.core_properties.created else "",
            "modified": str(doc.core_properties.modified) if doc.core_properties.modified else "",
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
        }

        # --- 提取所有段落并按 PARAGRAPHS_PER_PAGE 分页 ---
        all_paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables_data_all: list[list[list[str]]] = []

        # 提取所有表格
        for table in doc.tables:
            table_rows: list[list[str]] = []
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                table_rows.append(row_cells)
            if table_rows:
                tables_data_all.append(table_rows)

        # 分页：将段落分块
        page_num = 1
        for i in range(0, len(all_paragraphs), self.PARAGRAPHS_PER_PAGE):
            chunk_paragraphs = all_paragraphs[i : i + self.PARAGRAPHS_PER_PAGE]
            page_text = "\n".join(chunk_paragraphs)

            # 分配表格到页面（按顺序大致分配）
            page_tables: list[list[list[str]]] = []
            table_chunk_size = max(1, len(tables_data_all) // max(1, len(all_paragraphs) // self.PARAGRAPHS_PER_PAGE))
            table_start = (page_num - 1) * table_chunk_size
            table_end = table_start + table_chunk_size
            if table_start < len(tables_data_all):
                page_tables = tables_data_all[table_start:table_end]

            pages.append(
                PageContent(
                    page_num=page_num,
                    text=page_text,
                    tables=page_tables,
                )
            )
            page_num += 1

        # --- 提取图片 ---
        if self.extract_images:
            try:
                seen_hashes: set[str] = set()

                # 遍历文档所有关系和内联形状提取图片
                for rel in doc.part.rels.values():
                    if "image" in rel.reltype:
                        image_data = rel.target_part.blob
                        img_hash = hashlib.sha256(image_data).hexdigest()[:16]

                        if img_hash in seen_hashes:
                            continue
                        seen_hashes.add(img_hash)

                        # 确定图片所在的大致页码
                        img_page_num = min(page_num - 1, 1)

                        images.append(
                            ImageContent(
                                page_num=img_page_num,
                                image_data=image_data,
                                image_hash=img_hash,
                                image_ext="png",
                            )
                        )

                # 也从内联形状中提取
                for para in doc.paragraphs:
                    for run in para.runs:
                        if run._element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"):
                            # 通过 XML 提取嵌入的图片
                            for blip in run._element.iter(
                                "{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
                            ):
                                embed_id = blip.get(
                                    "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                                )
                                if embed_id:
                                    rel = doc.part.rels[embed_id]
                                    if rel and "image" in rel.reltype:
                                        image_data = rel.target_part.blob
                                        img_hash = hashlib.sha256(image_data).hexdigest()[:16]
                                        if img_hash not in seen_hashes:
                                            seen_hashes.add(img_hash)
                                            images.append(
                                                ImageContent(
                                                    page_num=1,
                                                    image_data=image_data,
                                                    image_hash=img_hash,
                                                    image_ext="png",
                                                )
                                            )

            except Exception as exc:
                logger.warning(f"DOCX 图片提取失败: {exc!s}")

        logger.info(
            f"DOCX 解析完成: {file_name}, "
            f"共 {len(pages)} 页（模拟分页）, {len(images)} 张图片, "
            f"{len(tables_data_all)} 个表格"
        )

        return ParseResult(
            file_name=file_name,
            file_type="docx",
            pages=pages,
            images=images,
            metadata=metadata,
        )
