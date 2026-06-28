"""
pytest 共享 fixture — BASS 测试基础设施。

提供内存 SQLite 数据库、样本文档、测试客户端，供所有测试模块复用。

注意：pytest-asyncio 1.3.0 要求异步 fixture 使用 @pytest_asyncio.fixture，
且同步 fixture 不能直接依赖异步 fixture。因此 db_session_factory 改为函数级
异步 fixture，在每次使用时动态创建引擎。
"""

from __future__ import annotations

import asyncio
import os
import uuid
from collections.abc import AsyncGenerator
from typing import Any

import pytest
import pytest_asyncio
from sqlalchemy.ext.asyncio import AsyncSession, async_sessionmaker, create_async_engine

from app.models.base import Base
from app.models.analysis import AnalysisTask
from app.models.project import BidDocument, Project


# ── helpers ──────────────────────────────────────────────────────────

def _make_uuid() -> str:
    return str(uuid.uuid4())


def _build_test_docx_bytes(
    author: str = "张三",
    company: str = "测试公司A",
    headings: list[str] | None = None,
    table_data: list[list[str]] | None = None,
) -> bytes:
    """用 python-docx 生成最小 .docx（内存字节）。"""
    from io import BytesIO
    from docx import Document
    from docx.shared import Inches

    doc = Document()

    cp = doc.core_properties
    cp.author = author
    cp.last_modified_by = author
    cp.title = f"投标文件 - {company}"
    cp.subject = "投标"
    cp.category = company

    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    if headings is None:
        headings = [
            "第一章 概述", "1.1 项目背景", "1.2 建设目标",
            "第二章 技术方案", "2.1 系统架构", "2.2 关键技术",
        ]
    for h in headings:
        doc.add_heading(h, level=1 if h.startswith("第") else 2)

    for i in range(3):
        doc.add_paragraph(f"这是第{i + 1}段正文内容，用于模拟投标文件的技术方案描述。")

    if table_data is None:
        table_data = [
            ["序号", "姓名", "职务", "职称"],
            ["1", "张三", "项目经理", "高级工程师"],
            ["2", "李四", "技术负责人", "工程师"],
        ]
    if table_data:
        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        for r, row in enumerate(table_data):
            for c, cell in enumerate(row):
                table.rows[r].cells[c].text = cell

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_test_pdf_bytes(
    author: str = "李四",
    title: str = "投标文件 - 测试公司",
    pages: int = 2,
) -> bytes:
    """用 fpdf2 生成最小 PDF（内存字节）。"""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_author(author)
    pdf.set_title(title)
    pdf.set_creator("Microsoft Word")
    pdf.set_producer("Adobe PDF Library 15.0")

    for p in range(pages):
        pdf.add_page()
        pdf.set_font("Helvetica", size=12)
        pdf.cell(200, 10, text=f"第{p + 1}页 投标文件内容", align="C")
        pdf.ln(20)
        pdf.set_font("Helvetica", size=10)
        pdf.multi_cell(0, 8, text=f"这是第{p + 1}页正文内容，模拟投标文件的技术方案描述。")

    return pdf.output()


# ── 异步 DB fixture（每个 test function 独立） ────────────────────────

@pytest_asyncio.fixture
async def db_session_factory():
    """创建内存 SQLite async_sessionmaker（函数级隔离）。"""
    engine = create_async_engine("sqlite+aiosqlite://", echo=False)

    async with engine.begin() as conn:
        await conn.run_sync(Base.metadata.create_all)

    factory = async_sessionmaker(engine, class_=AsyncSession, expire_on_commit=False)

    yield factory

    async with engine.begin() as conn:
        await conn.run_sync(Base.metadata.drop_all)
    await engine.dispose()


# ── 业务 fixture（每个 test function 独立） ───────────────────────────

@pytest_asyncio.fixture
async def sample_project(db_session_factory) -> Project:
    """创建一个带名称的示例项目。"""
    async with db_session_factory() as db:
        proj = Project(
            id=_make_uuid(), name="测试项目-TDD",
            description="TDD 单元测试项目", status="active",
        )
        db.add(proj)
        await db.commit()
        await db.refresh(proj)
        return proj


@pytest_asyncio.fixture
async def sample_documents(db_session_factory, sample_project: Project):
    """创建两个已完成解析的 BidDocument（含 content_text / file_metadata / extracted_tables）。"""
    p_id = sample_project.id
    doc1 = BidDocument(
        id=_make_uuid(), project_id=p_id,
        filename="标书A.docx", file_path=f"/tmp/{p_id}/标书A.docx",
        file_size=20480, file_type="docx",
        status="parsed", parse_status="completed",
        content_text=(
            "第一章 概述\n1.1 项目背景\n这是正文段落。\n"
            "第二章 技术方案\n2.1 系统架构\n2.2 关键技术"
        ),
        file_metadata={
            "author": "张三", "creator": "Microsoft Word",
            "producer": "Microsoft Word", "title": "投标文件 - 测试公司A",
            "company": "测试公司A", "last_modified_by": "张三",
        },
        extracted_tables=[{
            "page": 3, "rows": 3, "cols": 4,
            "header": ["序号", "姓名", "职务", "职称"],
            "data": [
                ["1", "张三", "项目经理", "高级工程师"],
                ["2", "李四", "技术负责人", "工程师"],
            ],
        }],
        page_count=15,
    )
    doc2 = BidDocument(
        id=_make_uuid(), project_id=p_id,
        filename="标书B.docx", file_path=f"/tmp/{p_id}/标书B.docx",
        file_size=18432, file_type="docx",
        status="parsed", parse_status="completed",
        content_text=(
            "第一章 概述\n1.1 项目背景\n这是略有不同的段落。\n"
            "第二章 技术方案\n2.1 总体设计\n2.2 核心功能"
        ),
        file_metadata={
            "author": "张三",  # ← 同一作者（元数据异常）
            "creator": "Microsoft Word", "producer": "Microsoft Word",
            "title": "投标文件 - 测试公司B", "company": "测试公司B",
            "last_modified_by": "张三",
        },
        extracted_tables=[{
            "page": 4, "rows": 3, "cols": 4,
            "header": ["序号", "姓名", "职务", "职称"],
            "data": [
                ["1", "王五", "项目经理", "高级工程师"],
                ["2", "赵六", "技术负责人", "助理工程师"],
            ],
        }],
        page_count=12,
    )
    async with db_session_factory() as db:
        db.add_all([doc1, doc2])
        await db.commit()
        await db.refresh(doc1)
        await db.refresh(doc2)
    return doc1, doc2


@pytest_asyncio.fixture
async def sample_analysis_task(db_session_factory, sample_project: Project):
    """创建一个 pending 状态的分析任务。"""
    async with db_session_factory() as db:
        task = AnalysisTask(
            id=_make_uuid(), project_id=sample_project.id,
            status="pending", task_type="full_analysis",
        )
        db.add(task)
        await db.commit()
        await db.refresh(task)
        return task


# ── 文件字节 fixtures（session 级别，只生成一次） ─────────────────────

@pytest.fixture(scope="session")
def docx_bytes_template_a() -> bytes:
    """模板A DOCX：作者张三、公司A。"""
    return _build_test_docx_bytes(author="张三", company="测试公司A")


@pytest.fixture(scope="session")
def docx_bytes_template_b() -> bytes:
    """模板B DOCX：作者张三（同一人）、公司B、不同布局。"""
    from io import BytesIO
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    cp = doc.core_properties
    cp.author = "张三"
    cp.last_modified_by = "张三"
    cp.title = "投标文件 - 测试公司B"
    cp.category = "测试公司B"

    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.5)   # ← 不同
        section.right_margin = Inches(1.0)

    for h in ["第一章 概述", "1.1 项目背景", "第二章 技术方案", "2.1 总体设计"]:
        doc.add_heading(h, level=1 if h.startswith("第") else 2)
    for i in range(3):
        doc.add_paragraph(f"这是公司B的第{i + 1}段正文内容。")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture(scope="session")
def pdf_bytes_doc_a() -> bytes:
    """PDF 样本A。"""
    return _build_test_pdf_bytes(author="李四", title="投标文件 - 测试公司A")


@pytest.fixture(scope="session")
def pdf_bytes_doc_b() -> bytes:
    """PDF 样本B。"""
    return _build_test_pdf_bytes(author="王五", title="投标文件 - 测试公司B")


# ── 集成测试 URL ─────────────────────────────────────────────────────

@pytest.fixture(scope="module")
def api_base_url() -> str:
    """集成测试使用的 API 基地址。"""
    return os.environ.get("BASS_TEST_API_URL", "http://localhost:8006/api/v1")
